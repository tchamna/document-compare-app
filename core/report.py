"""Generate Word (.docx) and PDF diff reports with bold-highlighted differences."""

from __future__ import annotations

import os
from pathlib import Path
from typing import List

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from core.helpers import LineDiff, word_diff_pairs


# ------------------------------------------------------------------ #
#  Register a Unicode-capable font (Charis SIL → Noto Sans fallback)
# ------------------------------------------------------------------ #
_FONT_REGISTERED = False
_FONT_NAME = "Helvetica"          # final fallback
_FONT_NAME_BOLD = "Helvetica-Bold"


def _register_unicode_font() -> None:
    """Register the best available Unicode TTF font with ReportLab."""
    global _FONT_REGISTERED, _FONT_NAME, _FONT_NAME_BOLD
    if _FONT_REGISTERED:
        return

    _FONT_REGISTERED = True

    # Candidate fonts in preference order: (name, regular path, bold path)
    win_fonts = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
    local_fonts = os.path.join(
        os.environ.get("LOCALAPPDATA", ""), "Microsoft", "Windows", "Fonts"
    )
    candidates = [
        (
            "CharisSIL",
            os.path.join(local_fonts, "CharisSIL-R.ttf"),
            os.path.join(local_fonts, "CharisSIL-B.ttf"),
        ),
        (
            "CharisSIL",
            os.path.join(win_fonts, "CharisSIL-R.ttf"),
            os.path.join(win_fonts, "CharisSIL-B.ttf"),
        ),
        (
            "NotoSans",
            os.path.join(win_fonts, "NotoSans-Regular.ttf"),
            os.path.join(win_fonts, "NotoSans-Bold.ttf"),
        ),
    ]

    for name, regular, bold in candidates:
        if os.path.isfile(regular):
            try:
                pdfmetrics.registerFont(TTFont(name, regular))
                bold_name = f"{name}-Bold"
                if os.path.isfile(bold):
                    pdfmetrics.registerFont(TTFont(bold_name, bold))
                    pdfmetrics.registerFontFamily(
                        name, normal=name, bold=bold_name,
                    )
                _FONT_NAME = name
                _FONT_NAME_BOLD = bold_name if os.path.isfile(bold) else name
                return
            except Exception:
                continue


def write_word_report(
    diffs: List[LineDiff],
    out_path: Path,
    original_name: str,
    corrige_name: str,
) -> None:
    """Write a Word document listing all differences.

    Differing words are **bolded** for easy visual scanning.
    """
    suffix = Path(original_name).suffix.lower()
    section_label = "Slide" if suffix == ".pptx" else "Paragraphe"

    doc = Document()
    doc.add_heading(
        "Differences between the original and corrected files", level=1
    )
    doc.add_paragraph("Files compared:")
    doc.add_paragraph(f"- Original file: {original_name}")
    doc.add_paragraph(f"- Corrected file: {corrige_name}")

    if not diffs:
        doc.add_paragraph("\nNo text differences detected.")
        doc.save(str(out_path))
        return

    current_slide = None
    for d in diffs:
        if d.slide_no != current_slide:
            current_slide = d.slide_no
            if section_label == "Slide":
                doc.add_heading(f"{section_label} {current_slide}", level=2)

        # --- Original ---
        doc.add_paragraph("Original:", style=None)
        p_orig = doc.add_paragraph()
        if d.original:
            for word, is_diff in word_diff_pairs(d.original, d.corrige):
                run = p_orig.add_run(word)
                if is_diff:
                    run.bold = True
                p_orig.add_run(" ")
        else:
            p_orig.add_run("(empty)")

        # --- Corrected ---
        doc.add_paragraph("Corrected:", style=None)
        p_corr = doc.add_paragraph()
        if d.corrige:
            for word, is_diff in word_diff_pairs(d.corrige, d.original):
                run = p_corr.add_run(word)
                if is_diff:
                    run.bold = True
                p_corr.add_run(" ")
        else:
            p_corr.add_run("(empty)")

        doc.add_paragraph("")  # spacer

    doc.save(str(out_path))


def write_pdf_report(
    diffs: List[LineDiff],
    out_path: Path,
    original_name: str,
    corrige_name: str,
) -> None:
    """Write a PDF document listing all differences.

    Differing words are marked with yellow highlight for easy visual scanning.
    """
    _register_unicode_font()

    suffix = Path(original_name).suffix.lower()
    section_label = "Slide" if suffix == ".pptx" else "Paragraphe"

    story = []

    # Base style with the Unicode font
    base_style = ParagraphStyle(
        "BaseUnicode",
        fontName=_FONT_NAME,
        fontSize=11,
        leading=14,
    )

    # Title
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=base_style,
        fontName=_FONT_NAME_BOLD,
        fontSize=18,
        textColor="#000000",
        spaceAfter=12,
        alignment=TA_CENTER,
    )
    story.append(
        Paragraph("Differences between the original and corrected files", title_style)
    )

    # File info
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(f"<b>Original file:</b> {original_name}", base_style))
    story.append(Paragraph(f"<b>Corrected file:</b> {corrige_name}", base_style))
    story.append(Spacer(1, 0.3 * inch))

    if not diffs:
        story.append(Paragraph("No text differences detected.", base_style))
    else:
        current_slide = None
        for d in diffs:
            # Slide header
            if d.slide_no != current_slide:
                current_slide = d.slide_no
                if section_label == "Slide":
                    heading_style = ParagraphStyle(
                        "CustomHeading",
                        parent=base_style,
                        fontName=_FONT_NAME_BOLD,
                        fontSize=14,
                        textColor="#000000",
                        spaceAfter=6,
                        spaceBefore=12,
                    )
                    story.append(
                        Paragraph(f"{section_label} {current_slide}", heading_style)
                    )

            # Original
            story.append(
                Paragraph("<b>Original:</b>", base_style)
            )
            orig_html = _words_to_html(d.original, d.corrige)
            story.append(Paragraph(orig_html, base_style))

            # Corrected
            story.append(Spacer(1, 0.1 * inch))
            story.append(
                Paragraph("<b>Corrected:</b>", base_style)
            )
            corr_html = _words_to_html(d.corrige, d.original)
            story.append(Paragraph(corr_html, base_style))

            story.append(Spacer(1, 0.2 * inch))

    # Generate PDF
    doc = SimpleDocTemplate(str(out_path), pagesize=letter)
    doc.build(story)


def _words_to_html(text: str, other: str) -> str:
    """Convert text with differences to HTML with yellow highlight (ReportLab compatible).
    
    Highlights only words that differ, using intelligent word-level alignment.
    """
    if not text:
        return '<i>(empty)</i>'
    
    from difflib import SequenceMatcher
    
    words_text = text.split()
    words_other = other.split()
    
    # Use SequenceMatcher to align words intelligently
    sm = SequenceMatcher(None, words_other, words_text)
    
    # Track which words in text are different
    different_indices = set()
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag != 'equal':
            # Mark these words as different
            for idx in range(j1, j2):
                different_indices.add(idx)
    
    # Build HTML with highlighting for different words
    parts = []
    for idx, word in enumerate(words_text):
        if idx in different_indices:
            parts.append(f'<font backColor="#ffd54f"><b>{word}</b></font>')
        else:
            parts.append(word)
    
    return " ".join(parts)
